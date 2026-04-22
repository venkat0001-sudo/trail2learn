"""
DOCX Parser — Extracts feature entries from the OEM specification document.

This module implements the two-pass extraction strategy we locked in during
design. It walks the DOCX in document order (via XML iteration), maintains
a heading hierarchy stack, and for each table it encounters decides whether
to ignore it (bit-field spec) or expand it (feature listing).

WHY WE ITERATE XML DIRECTLY INSTEAD OF doc.paragraphs / doc.tables:
  python-docx's high-level accessors split paragraphs and tables into
  separate lists, losing document order. We need order to link each table
  to the heading that precedes it. The underlying element tree preserves
  order, so we walk it manually.

OUTPUT:
  A list of RawFeatureEntry (see pydantic_schemas.py) plus a SectionIndex
  that maps section numbers to positions in the flat text, used later for
  cross-reference resolution and Qwen context assembly.
"""

from __future__ import annotations

import re
import logging
from dataclasses import dataclass, field
from typing import Optional
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from pydantic_schemas import FeatureKey, RawFeatureEntry


logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────
# Bit-field header detection
#
# When a table's first column header matches any of these patterns, we
# classify it as a bit-field specification table and do NOT emit rows as
# features. Derived from your actual spec screenshots (Image 12, 13).
# ─────────────────────────────────────────────────────────────────────────

BITFIELD_HEADER_TOKENS = {
    # Data width descriptors
    "dword", "word", "byte", "bytes",
    # Bit position descriptors
    "bit", "bits", "bit position", "bit range",
    # Field-layout descriptors
    "offset", "offset start", "offset end", "field", "register",
    # Value tables (Identify Controller sub-tables)
    "value", "size", "size (bytes)",
    # Measurement parameter tables (shock/vibration)
    "half-sine pulse duration", "pulse peak level spec",
    "half-sine pulse duration (ms)", "pulse peak level spec (g)",
    # Form factor sub-tables
    "form factor", "max weight", "form factor description",
}


def is_bitfield_header(text: str) -> bool:
    """
    Return True if the given cell text looks like a bit-field table header.
    Case-insensitive, whitespace-tolerant.
    """
    if not text:
        return False
    normalized = " ".join(text.strip().lower().split())
    if normalized in BITFIELD_HEADER_TOKENS:
        return True
    # Also catch headers that start with these tokens followed by more text
    for token in BITFIELD_HEADER_TOKENS:
        if normalized.startswith(token + " ") or normalized.startswith(token + ":"):
            return True
    return False


# ─────────────────────────────────────────────────────────────────────────
# Section number regex
#
# NVMe / OEM specs use standard hierarchical numbering: X, X.Y, X.Y.Z, ...
# Optionally followed by a title. This regex extracts both.
# ─────────────────────────────────────────────────────────────────────────

# Matches leading "4", "4.1", "4.1.2", "4.1.2.3" with optional trailing title.
SECTION_NUMBER_RE = re.compile(
    r"^\s*(?P<num>\d{1,2}(?:\.\d{1,3}){0,4})\s*(?P<title>.*?)\s*$"
)


def parse_heading_text(text: str) -> tuple[Optional[str], str]:
    """
    Split heading text into (section_number, title).
    If no number is found, returns (None, full_text).

    Examples:
        "4.8.3 SCP Feature Enable/Disable" → ("4.8.3", "SCP Feature Enable/Disable")
        "Mechanical"                        → (None, "Mechanical")
        "4. Behavioral Specification"       → ("4", "Behavioral Specification")
    """
    if not text:
        return None, ""
    text = text.strip()
    match = SECTION_NUMBER_RE.match(text)
    if match:
        return match.group("num"), match.group("title").strip()
    return None, text


def parent_of(section_number: str) -> Optional[str]:
    """
    Compute the parent section number by stripping the last dotted segment.

    Examples:
        "4.8.3" → "4.8"
        "4.8"   → "4"
        "4"     → None
    """
    if not section_number or "." not in section_number:
        return None
    return section_number.rsplit(".", 1)[0]


# ─────────────────────────────────────────────────────────────────────────
# HeadingStack — maintains the top-chapter context as we walk the document
# ─────────────────────────────────────────────────────────────────────────

@dataclass
class HeadingContext:
    """The current heading hierarchy at any point in the document walk."""
    # Keyed by heading level (1, 2, 3, ...), stores (section_number, title) pairs.
    # When we see a Heading N, we clear all levels >= N and set level N.
    stack: dict[int, tuple[Optional[str], str]] = field(default_factory=dict)

    def observe_heading(self, level: int, section_number: Optional[str], title: str):
        """Record a new heading at the given level, clearing deeper levels."""
        # Clear any levels deeper than the new one — they're now out of scope.
        deeper_levels = [lvl for lvl in self.stack if lvl >= level]
        for lvl in deeper_levels:
            del self.stack[lvl]
        self.stack[level] = (section_number, title)

    def top_chapter_label(self) -> Optional[str]:
        """
        Return the Heading 1 label formatted for the Section column,
        e.g., "4. Behavioral Specification".

        If Heading 1 has a section number like "4" and title "Behavioral
        Specification", we combine them as "4. Behavioral Specification".
        """
        if 1 not in self.stack:
            return None
        num, title = self.stack[1]
        if num:
            return f"{num}. {title}"
        return title

    def current_heading(self) -> Optional[tuple[Optional[str], str, int]]:
        """Return (section_number, title, level) of the deepest current heading."""
        if not self.stack:
            return None
        deepest_level = max(self.stack.keys())
        num, title = self.stack[deepest_level]
        return (num, title, deepest_level)


# ─────────────────────────────────────────────────────────────────────────
# Heading level detection from DOCX style
#
# python-docx exposes paragraph.style.name. Standard Word heading styles
# are named 'Heading 1', 'Heading 2', etc. Some docs use custom style
# names — we detect those by looking for 'heading' substring.
# ─────────────────────────────────────────────────────────────────────────

def heading_level(paragraph: Paragraph) -> Optional[int]:
    """
    Return the heading level (1-9) of a paragraph, or None if not a heading.
    """
    style_name = paragraph.style.name if paragraph.style else ""
    if not style_name:
        return None
    style_lower = style_name.lower()

    # Standard "Heading N" styles
    match = re.match(r"heading\s*(\d+)", style_lower)
    if match:
        level = int(match.group(1))
        if 1 <= level <= 9:
            return level

    # "Title" is often used as level 1
    if style_lower == "title":
        return 1

    return None


# ─────────────────────────────────────────────────────────────────────────
# Document order iteration
#
# We walk the doc.element.body children. Each child is either a <w:p>
# (paragraph) or a <w:tbl> (table). We yield them as python-docx objects.
# ─────────────────────────────────────────────────────────────────────────

def iter_document_elements(doc: DocxDocument):
    """
    Yield each top-level paragraph or table in document order.

    Returns tuples (kind, element) where kind is 'paragraph' or 'table'
    and element is a Paragraph or Table object.
    """
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            yield ("paragraph", Paragraph(child, doc))
        elif tag == qn("w:tbl"):
            yield ("table", Table(child, doc))
        # Ignore other element types (section properties, etc.)


# ─────────────────────────────────────────────────────────────────────────
# Table handling
#
# For each table we classify as feature-listing or bit-field.
# Feature-listing tables produce one RawFeatureEntry per row.
# ─────────────────────────────────────────────────────────────────────────

def get_cell_text(cell: _Cell) -> str:
    """
    Return all text in a cell, joined with spaces. python-docx separates
    cell content into paragraphs; we flatten them.
    """
    parts = []
    for para in cell.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    return " ".join(parts)


def classify_table(table: Table) -> str:
    """
    Classify a table as 'bitfield', 'feature_list', or 'unknown'.

    Rules (from design FQ-1):
      - First column header matches a known bit-field header token → 'bitfield'
      - Otherwise if first column cells look like feature names → 'feature_list'
      - If we can't tell → 'unknown' (we skip to be safe)
    """
    if not table.rows:
        return "unknown"

    header_row = table.rows[0]
    if not header_row.cells:
        return "unknown"

    first_header = get_cell_text(header_row.cells[0])

    if is_bitfield_header(first_header):
        return "bitfield"

    # Check if ANY column header is a bitfield token. If so, likely a
    # bit-field table with a blank first column.
    for cell in header_row.cells:
        if is_bitfield_header(get_cell_text(cell)):
            return "bitfield"

    # Heuristic: feature-listing tables usually have 2 columns (name, description)
    # and the first column has descriptive feature names (mostly alphabetic).
    # Bit-field tables often have 3-4+ columns.
    if len(header_row.cells) >= 2:
        # Look at a data row (skip header)
        if len(table.rows) >= 2:
            first_data_cell = get_cell_text(table.rows[1].cells[0])
            # Feature names are mostly alphabetic; bit-field values are numeric/hex
            if first_data_cell and any(c.isalpha() for c in first_data_cell):
                # And not dominated by digits (bit ranges like "31:16")
                alpha_ratio = sum(c.isalpha() for c in first_data_cell) / max(len(first_data_cell), 1)
                if alpha_ratio > 0.5:
                    return "feature_list"

    return "unknown"


def extract_feature_rows_from_table(
    table: Table,
    parent_section: Optional[str],
    top_chapter: Optional[str],
    parent_full_text: str,
) -> list[RawFeatureEntry]:
    """
    For a feature-listing table, emit one RawFeatureEntry per data row.
    parent_section is the section number of the heading that contains this
    table (e.g., "4.7" for the Firmware Features table).
    """
    entries = []
    rows = table.rows
    if len(rows) < 2:
        return entries

    # Assume row 0 is the header; skip it.
    for row in rows[1:]:
        if not row.cells:
            continue
        topic = get_cell_text(row.cells[0])
        if not topic:
            continue

        # Grab content of the second column if present — that's the
        # per-row requirement text (e.g., "SSD shall comply with SMART...").
        cell_text = ""
        if len(row.cells) >= 2:
            cell_text = get_cell_text(row.cells[1])

        key = FeatureKey(origin="table_row", parent=parent_section, topic=topic)

        entries.append(RawFeatureEntry(
            key=key,
            topic=topic,
            section_number=None,
            parent_section=parent_section,
            top_chapter=top_chapter,
            full_text=parent_full_text,   # for LLM context
            cell_text=cell_text,          # the specific requirement text
        ))

    return entries


# ─────────────────────────────────────────────────────────────────────────
# Section text accumulator
#
# For each heading we emit, we want its 'full_text' field to contain all
# paragraphs (and prose from any non-feature tables) that appear under it
# until the NEXT heading at the same or higher level.
#
# We accomplish this with a small state machine: we track the current
# heading's accumulated body text; when a new heading arrives, we flush
# the previous section's entry with its accumulated text.
# ─────────────────────────────────────────────────────────────────────────

@dataclass
class PendingSection:
    """A section whose text is still being accumulated."""
    section_number: Optional[str]
    title: str
    level: int
    parent_section: Optional[str]
    top_chapter: Optional[str]
    body_parts: list[str] = field(default_factory=list)
    # Table-row entries found UNDER this section (after its heading,
    # before the next heading). These inherit this section as parent.
    child_table_rows: list[RawFeatureEntry] = field(default_factory=list)

    def add_text(self, text: str):
        if text and text.strip():
            self.body_parts.append(text.strip())

    def full_text(self) -> str:
        return "\n\n".join(self.body_parts)

    def to_entry(self) -> RawFeatureEntry:
        topic = (f"{self.section_number} {self.title}".strip()
                 if self.section_number else self.title)
        key = FeatureKey(origin="section", section=self.section_number)
        return RawFeatureEntry(
            key=key,
            topic=topic,
            section_number=self.section_number,
            parent_section=self.parent_section,
            top_chapter=self.top_chapter,
            full_text=self.full_text(),
            cell_text=None,
        )


# ─────────────────────────────────────────────────────────────────────────
# Main parser entry point
# ─────────────────────────────────────────────────────────────────────────

def parse_oem_docx(docx_path: Path) -> list[RawFeatureEntry]:
    """
    Parse an OEM specification DOCX and return all detected feature entries.

    This is the single public entry point of this module. Call it with the
    path to an OEM .docx file and get back a list of RawFeatureEntry objects
    ready to feed into Stage 1.

    Rules applied:
      - Every DOCX heading becomes a section-origin feature entry.
      - Feature-listing tables (first-column names) produce table-row entries,
        one per data row, with parent_section set to the enclosing section.
      - Bit-field tables are skipped (the parent heading covers the feature).
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(str(docx_path))
    heading_ctx = HeadingContext()
    entries: list[RawFeatureEntry] = []

    # The section we're currently accumulating body text for.
    pending: Optional[PendingSection] = None

    def flush_pending():
        """Emit the pending section's entry plus any child table rows."""
        nonlocal pending
        if pending is None:
            return
        # Emit the section-origin entry (always, regardless of whether it had children)
        section_entry = pending.to_entry()
        entries.append(section_entry)
        # Emit any table-row children. We update their full_text AFTER we
        # know the final accumulated body text of this section.
        for child in pending.child_table_rows:
            child.full_text = section_entry.full_text
            entries.append(child)
        pending = None

    for kind, elem in iter_document_elements(doc):

        if kind == "paragraph":
            text = elem.text
            lvl = heading_level(elem)

            if lvl is not None:
                # It's a heading — flush the previous section, start a new one.
                flush_pending()

                section_num, title = parse_heading_text(text)
                heading_ctx.observe_heading(lvl, section_num, title)

                # Compute parent section: for a heading, parent is either the
                # numeric parent (4.8.3's parent is 4.8) or, if no number, None.
                parent_num = parent_of(section_num) if section_num else None
                top_chapter = heading_ctx.top_chapter_label()

                pending = PendingSection(
                    section_number=section_num,
                    title=title,
                    level=lvl,
                    parent_section=parent_num,
                    top_chapter=top_chapter,
                )
            else:
                # Regular paragraph — append to the current section's body.
                if pending is not None and text.strip():
                    pending.add_text(text)

        elif kind == "table":
            # Classify and act accordingly.
            table_class = classify_table(elem)

            if table_class == "feature_list":
                if pending is None:
                    # Rare: a table with no preceding heading. Skip with warning.
                    logger.warning("Feature-listing table found with no parent heading; skipping")
                    continue
                # Emit row entries under the current section.
                # Use whatever text we've accumulated so far as parent_full_text.
                parent_full_text = pending.full_text()
                row_entries = extract_feature_rows_from_table(
                    elem,
                    parent_section=pending.section_number,
                    top_chapter=pending.top_chapter,
                    parent_full_text=parent_full_text,
                )
                pending.child_table_rows.extend(row_entries)
            else:
                # Bit-field or unknown — fold the prose content of the table
                # (if any) into the current section's body text. This preserves
                # content like "Admin Opcode: 09h  Feature ID: D8h" that appears
                # above bit-field tables and isn't a separate paragraph.
                if pending is not None:
                    for row in elem.rows:
                        for cell in row.cells:
                            cell_text = get_cell_text(cell)
                            if cell_text:
                                pending.add_text(cell_text)

    # Flush the final section.
    flush_pending()

    logger.info(f"Parsed {len(entries)} feature entries from {docx_path.name}")
    return entries
